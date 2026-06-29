"""
Generate a Consumer Commission complaint:
  - "polished_narrative" pass via Gemini (turns the user's intake into formal legal prose)
  - DOCX assembly via python-docx
  - PDF assembly via reportlab

All sections, reliefs, and party blocks come from the structured intake — Gemini
only polishes the narrative paragraphs, never invents facts.
"""
from __future__ import annotations

import time
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from . import config
from .intake import ComplaintIntake
from .jurisdiction import JurisdictionVerdict
from .knowledge import category, section_by_id
from .llm import gemini_generate


# ── Per-category statutory vocabulary injected into the narrative prompt ──────
_CATEGORY_LEGAL_VOCAB: dict[str, str] = {
    "defective_goods": (
        "This matter involves a defective product. Use the term 'defect' as defined under "
        "Section 2(10) of the Consumer Protection Act, 2019 (meaning any fault, imperfection, "
        "shortcoming or inadequacy in quality, quantity, potency, purity or standard). "
        "Reference 'product liability' under Section 2(34) where the seller/manufacturer "
        "is responsible for harm caused by the defective product. Where the Opposite Party "
        "refused to repair or replace, also characterise this as 'deficiency in service' "
        "u/s 2(11) of the Act."
    ),
    "deficient_service": (
        "This matter involves deficiency in service. Use the precise statutory phrase "
        "'deficiency in service' as defined under Section 2(11) of the Consumer Protection "
        "Act, 2019 (meaning any fault, imperfection, shortcoming or inadequacy in the "
        "quality, nature and manner of performance of any service). Also reference Section "
        "2(42) for the definition of 'service'. Where the Opposite Party charged a fee but "
        "failed to render service of contracted quality, that constitutes a deficiency."
    ),
    "unfair_trade_practice": (
        "This matter involves an unfair trade practice. Use the precise phrase 'unfair trade "
        "practice' as defined under Section 2(47) of the Consumer Protection Act, 2019. "
        "Identify the specific sub-clause that applies — e.g., false representation of "
        "quality/standard (Section 2(47)(i)), misleading advertisement (Section 2(47)(vi)), "
        "bait-and-switch offers (Section 2(47)(viii)). Characterise the Opposite Party's "
        "conduct as a 'false and misleading representation made with intent to induce the "
        "Complainant to enter into the transaction'."
    ),
    "medical_negligence": (
        "This matter involves medical negligence constituting deficiency in service. Use "
        "'deficiency in service' as defined under Section 2(11) of the Consumer Protection "
        "Act, 2019. Establish that the Complainant is a 'consumer' within Section 2(7) of "
        "the Act as the medical services were availed for consideration. Describe the "
        "standard of care that was owed and the negligent departure therefrom. Reference "
        "that the hospital/clinic is a 'service provider' within the meaning of the Act."
    ),
    "real_estate": (
        "This matter involves a builder-buyer dispute constituting deficiency in service "
        "and/or unfair trade practice. Use 'deficiency in service' u/s 2(11) of the Act "
        "for delayed or defective possession. Reference 'unfair trade practice' u/s 2(47) "
        "for any false promises regarding amenities, completion dates, or specifications. "
        "If the project is RERA-registered, note the registration number and that the "
        "Consumer Commission's jurisdiction is in addition to RERA, as per Section 100 of "
        "the Consumer Protection Act, 2019."
    ),
    "e_commerce": (
        "This matter involves an e-commerce / online marketplace dispute. Reference Section "
        "94 of the Consumer Protection Act, 2019 which imposes obligations on e-commerce "
        "entities. Use 'deficiency in service' u/s 2(11) and 'defect' u/s 2(10) as "
        "applicable. Note that both the marketplace platform and the actual seller may be "
        "jointly and severally liable for the grievance. Reference the Consumer Protection "
        "(E-Commerce) Rules, 2020 if the platform failed to comply with its mandatory "
        "disclosure or grievance redressal obligations."
    ),
}

_NARRATIVE_SYSTEM = """\
You are a qualified advocate / senior paralegal drafting a formal complaint petition for \
filing before an Indian Consumer Disputes Redressal Commission under the Consumer \
Protection Act, 2019. You have extensive experience drafting Indian legal pleadings.

MANDATORY STYLE RULES — follow all of these without exception:

1. PARAGRAPH FORMAT: Number each paragraph sequentially. Every paragraph MUST begin \
   with the word "That". Use the format: "Para 1.  That..." / "Para 2.  That..." etc.

2. VOICE AND PERSON: Write entirely in formal third person.
   - Refer to the person filing → "the Complainant" (never "I", "we", or "you")
   - Refer to the other side → "the Opposite Party" / "Opposite Party No. 1" etc.
   - Refer to the court → "this Hon'ble Commission" or "this Hon'ble District/State \
     Consumer Disputes Redressal Commission"

3. LEGAL TERMINOLOGY: Use precise statutory and legal phrases throughout:
   - "deficiency in service" (not "bad service" or "poor service")
   - "defect in the said goods" (not "broken" or "damaged product")
   - "unfair trade practice" (not "misleading" or "false promises")
   - "the cause of action arose on..." (for the triggering date)
   - "inter alia" for "among other things"
   - "ex facie" for "on the face of it"
   - "humbly submits", "respectfully avers", "humbly states"
   - "the Opposite Party is estopped from..." (where applicable)
   - "without prejudice to the generality of the foregoing" (for additional points)

4. STATUTORY CITATIONS: Cite the Consumer Protection Act, 2019 inline:
   - "as defined under Section 2(10) of the Consumer Protection Act, 2019 \
     (hereinafter referred to as 'the Act')" — use this full form on first reference only
   - Thereafter use: "u/s 2(11) of the Act", "under Section 35 of the Act"

5. MONETARY NOTATION: Always write: "Rs. [amount]/- (Rupees [amount in words] Only)"
   Example: "Rs. 25,000/- (Rupees Twenty-Five Thousand Only)"

6. DATE NOTATION: Use the format "[DD.MM.YYYY]" or "on [date] (hereinafter 'the \
   Transaction Date')" for key dates.

7. EXHIBIT REFERENCES: Label all documents as "Annexure-[Letter]" in the order they \
   are first introduced. Example: "a copy of the invoice is annexed hereto and marked \
   as Annexure-A". Assign: invoice/receipt → Annexure-A; warranty card → Annexure-B; \
   correspondence → Annexure-C; pre-litigation notice → Annexure-D; etc.

8. STRICT FACTUAL ACCURACY: Use ONLY the facts supplied in the intake. Do NOT invent \
   any dates, names, amounts, documents, or facts not present in the intake data.

9. SCOPE: Do NOT include the prayer / relief paragraphs — those are drafted separately.

10. TONE: Formal, measured, and precise. No colloquialisms, contractions, or plain-\
    English summaries. Every sentence should read as if written by an experienced advocate.
"""


def _polish_narrative(intake: ComplaintIntake) -> str:
    cat = category(intake.issue.category_key)
    vocab_hint = _CATEGORY_LEGAL_VOCAB.get(intake.issue.category_key, "")

    txn_amt = intake.transaction.transaction_amount_inr

    facts_block = (
        f"COMPLAINANT: {intake.complainant.name}, residing at "
        f"{intake.complainant.address_line}, {intake.complainant.city}, "
        f"{intake.complainant.state} - {intake.complainant.pincode}.\n"
        f"OPPOSITE PARTY: {intake.opposite_party.name}, having its office / place of "
        f"business at {intake.opposite_party.address_line}, {intake.opposite_party.city}, "
        f"{intake.opposite_party.state} - {intake.opposite_party.pincode}.\n"
        f"COMPLAINT CATEGORY: {cat['label']}.\n"
        f"DATE OF TRANSACTION: {intake.transaction.transaction_date}.\n"
        f"DATE CAUSE OF ACTION AROSE: {intake.transaction.cause_of_action_date}.\n"
        f"AMOUNT PAID (INR): Rs. {txn_amt:,.0f}/- \n"
        f"INVOICE / ORDER REFERENCE: {intake.transaction.invoice_or_reference or 'N/A'}.\n"
        f"PAYMENT MODE: {intake.transaction.payment_mode or 'N/A'}.\n"
        f"PRE-LITIGATION LEGAL NOTICE SENT: "
        f"{'Yes' if intake.issue.pre_litigation_notice_sent else 'No'}.\n"
        f"RESPONSE TO NOTICE: "
        f"{intake.issue.notice_response_summary or 'No response received / ignored'}.\n"
        f"DOCUMENTARY EVIDENCE AVAILABLE: "
        f"{', '.join(intake.issue.evidence_available) or 'As detailed in complaint'}.\n\n"
        f"COMPLAINANT'S ACCOUNT OF FACTS (verbatim — convert to legal form):\n"
        f"{intake.issue.detailed_narrative}\n"
    )

    prompt = (
        "Draft the STATEMENT OF FACTS section of a Consumer Commission complaint petition "
        "under the Consumer Protection Act, 2019.\n\n"
        "Produce between 10 and 16 numbered paragraphs in the format "
        "'Para 1.  That...' / 'Para 2.  That...' etc.\n\n"
        "The paragraphs MUST cover, in this order:\n"
        "  (a) Para 1: The Complainant's status as a 'consumer' u/s 2(7) of the Act — "
        "what was purchased/contracted, for what purpose (personal/household use), "
        "and that it was not for resale or any commercial purpose.\n"
        "  (b) Para 2: The Opposite Party's status — who they are, their registered "
        "office, and their role as manufacturer / seller / service provider.\n"
        "  (c) Para 3–4: The transaction in detail — product/service, date of "
        "purchase/contract, amount paid (Rs. X/- format), invoice/reference number, "
        "mode of payment. Introduce the invoice as Annexure-A.\n"
        "  (d) Para 5: Representation made by the Opposite Party — what was promised "
        "or advertised (quality, warranty, delivery date, specifications).\n"
        "  (e) Para 6–8: The defect / deficiency / unfair practice — what went wrong, "
        "when it was discovered, what the deficiency / defect consists of (use statutory "
        "terminology), and how it falls within the definitions under the Act.\n"
        "  (f) Para 9–10: Steps taken to resolve — complaints lodged (with dates and "
        "reference numbers), pre-litigation legal notice sent (if any, as Annexure-D), "
        "and the Opposite Party's response or failure to respond.\n"
        "  (g) Para 11: The resulting harm — financial loss, mental agony, harassment, "
        "loss of time and incidental expenses suffered by the Complainant.\n"
        "  (h) Para 12: The cause of action — when it arose, that it is a continuing "
        "cause of action (if applicable), and that the complaint is filed within the "
        "limitation period prescribed under Section 69 of the Act.\n"
        "  (i) Para 13: Jurisdiction — pecuniary (claim amount and applicable forum "
        "under the 2021 Jurisdiction Rules) and territorial (complainant's place of "
        "residence / OP's place of business / place where cause of action arose).\n\n"
        f"CATEGORY-SPECIFIC STATUTORY VOCABULARY TO USE:\n{vocab_hint}\n\n"
        f"INTAKE DATA:\n{facts_block}\n\n"
        "EXHIBIT LABELLING: Assign Annexure-A to the invoice/receipt, Annexure-B to the "
        "warranty card (if any), Annexure-C to correspondence with the Opposite Party, "
        "Annexure-D to the pre-litigation notice (if sent), in the order first referenced."
    )

    return gemini_generate(
        prompt,
        system_instruction=_NARRATIVE_SYSTEM,
        temperature=0.15,
        max_output_tokens=2800,
    ).strip()


@dataclass
class DraftBundle:
    title: str
    polished_narrative: str
    relevant_sections: list[dict]
    cause_title: str
    docx_path: Path
    pdf_path: Path
    text: str


def _cause_title(intake: ComplaintIntake, verdict: JurisdictionVerdict) -> str:
    return (
        f"Consumer Complaint No. _____ of {date.today().year}\n"
        f"Before the Hon'ble {verdict.forum} Consumer Disputes Redressal Commission"
    )


def _relevant_sections(intake: ComplaintIntake) -> list[dict]:
    cat = category(intake.issue.category_key)
    out: list[dict] = []
    for sid in cat["primary_sections"]:
        s = section_by_id(sid)
        if s:
            out.append(s)
        else:
            out.append({"section": sid, "title": "(see CPA 2019)", "summary": ""})
    return out


def _inr_words(amount: float) -> str:
    a = int(round(amount))
    crore, rem = divmod(a, 10_000_000)
    lakh, rem = divmod(rem, 100_000)
    thousand, rem = divmod(rem, 1_000)
    parts = []
    if crore:
        parts.append(f"{crore} Crore")
    if lakh:
        parts.append(f"{lakh} Lakh")
    if thousand:
        parts.append(f"{thousand} Thousand")
    if rem or not parts:
        parts.append(str(rem))
    return "Rupees " + " ".join(parts) + " Only"


def _reliefs_block(intake: ComplaintIntake) -> list[str]:
    cat = category(intake.issue.category_key)
    items: list[str] = []

    txn_amt = intake.transaction.transaction_amount_inr
    txn_str = f"Rs. {txn_amt:,.0f}/- ({_inr_words(txn_amt)})"

    if intake.relief.refund_requested:
        items.append(
            f"Direct the Opposite Party to refund to the Complainant the sum of "
            f"{txn_str}, being the amount paid by the Complainant, together with "
            f"interest thereon at the rate of 9% per annum (or at such rate as this "
            f"Hon'ble Commission may deem fit and proper in the circumstances of the "
            f"case) from the date of payment till the date of actual realisation."
        )
    if intake.relief.replacement_requested:
        items.append(
            "Direct the Opposite Party to forthwith replace the defective "
            "product / deficient service with a defect-free product of the same "
            "description and specifications / render the contracted service in a "
            "proper, complete and satisfactory manner, at no additional cost to the "
            "Complainant, failing which direct the Opposite Party to refund the "
            "full consideration paid."
        )
    agony = intake.relief.compensation_for_mental_agony_inr
    if agony > 0:
        agony_str = f"Rs. {agony:,.0f}/- ({_inr_words(agony)})"
        items.append(
            f"Direct the Opposite Party to pay the Complainant a sum of {agony_str} "
            f"as compensation for the mental agony, harassment, inconvenience, "
            f"hardship, and incidental expenses suffered and incurred by the Complainant "
            f"as a direct and proximate consequence of the deficient and negligent "
            f"conduct of the Opposite Party."
        )
    if intake.relief.litigation_cost_requested:
        items.append(
            "Direct the Opposite Party to pay the costs of this complaint / litigation, "
            "as assessed by this Hon'ble Commission."
        )
    for c in intake.relief.custom_reliefs:
        if c.strip():
            items.append(c.strip())

    if not items:
        for r in cat["typical_reliefs"]:
            items.append(r)

    return items


def _grounds_section(intake: ComplaintIntake, sections: list[dict]) -> str:
    lines: list[str] = []
    lines.append("GROUNDS:")
    lines.append(
        "This complaint is maintainable and the Complainant is entitled to the reliefs "
        "prayed for on the following grounds, amongst others:"
    )

    roman = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X",
             "XI", "XII", "XIII", "XIV", "XV"]
    idx = 0

    for s in sections:
        r = roman[idx] if idx < len(roman) else str(idx + 1)
        idx += 1
        lines.append(
            f"  {r}. BECAUSE Section {s['section']} of the Consumer Protection Act, "
            f"2019 — '{s['title']}' — {s['summary'].rstrip('.')} — is squarely "
            f"attracted on the facts of the present case."
        )

    r = roman[idx] if idx < len(roman) else str(idx + 1)
    idx += 1
    lines.append(
        f"  {r}. BECAUSE the Complainant is a 'consumer' within the meaning of "
        f"Section 2(7) of the Consumer Protection Act, 2019, having purchased the "
        f"goods / availed the services from the Opposite Party for a valid "
        f"consideration, for personal / household use and not for any commercial "
        f"purpose or for resale."
    )

    r = roman[idx] if idx < len(roman) else str(idx + 1)
    idx += 1
    lines.append(
        f"  {r}. BECAUSE the complaint is filed within the period of limitation of "
        f"two years from the date on which the cause of action arose, as prescribed "
        f"under Section 69 of the Consumer Protection Act, 2019, and is therefore "
        f"within limitation."
    )

    r = roman[idx] if idx < len(roman) else str(idx + 1)
    idx += 1
    lines.append(
        f"  {r}. BECAUSE this Hon'ble Commission has pecuniary jurisdiction to "
        f"entertain this complaint in terms of the Consumer Protection (Jurisdiction "
        f"of the District Commission, the State Commission and the National Commission) "
        f"Rules, 2021 read with Sections 34, 47, and 58 of the Consumer Protection "
        f"Act, 2019, the aggregate claim being within the prescribed pecuniary limits."
    )

    r = roman[idx] if idx < len(roman) else str(idx + 1)
    lines.append(
        f"  {r}. BECAUSE this Hon'ble Commission has territorial jurisdiction to "
        f"entertain this complaint inasmuch as the Complainant resides / works for "
        f"gain within the jurisdiction of this Commission, as expressly permitted "
        f"under Section 34(2)(c) of the Consumer Protection Act, 2019."
    )

    return "\n".join(lines)


def _text_complaint(
    intake: ComplaintIntake,
    verdict: JurisdictionVerdict,
    narrative: str,
    sections: list[dict],
) -> str:
    parts: list[str] = []

    # Court header
    if verdict.forum == "District":
        court_line = (
            f"BEFORE THE HON'BLE {intake.complainant.city.upper()} "
            f"DISTRICT CONSUMER DISPUTES REDRESSAL COMMISSION,\n"
            f"{intake.complainant.state.upper()}"
        )
    elif verdict.forum == "State":
        court_line = (
            f"BEFORE THE HON'BLE STATE CONSUMER DISPUTES REDRESSAL COMMISSION,\n"
            f"{intake.complainant.state.upper()}"
        )
    else:
        court_line = (
            "BEFORE THE HON'BLE NATIONAL CONSUMER DISPUTES REDRESSAL COMMISSION,\n"
            "NEW DELHI"
        )

    parts.append(court_line)
    parts.append(f"Consumer Complaint No. _____ of {date.today().year}")
    parts.append("IN THE MATTER OF:")

    # Complainant block
    c = intake.complainant
    parts.append(
        f"{c.name}\n"
        f"S/o / D/o / W/o: _______________\n"
        f"{c.address_line}\n"
        f"{c.city}, {c.state} \u2013 {c.pincode}\n"
        f"Phone: {c.phone or '_______________'}\n"
        f"Email: {c.email or '_______________'}"
        f"\n\n"
        f"                                                               ...COMPLAINANT"
    )

    parts.append("VERSUS")

    # Opposite party block
    o = intake.opposite_party
    parts.append(
        f"{o.name}\n"
        f"{o.address_line}\n"
        f"{o.city}, {o.state} \u2013 {o.pincode}\n"
        f"Phone: {o.phone or '_______________'}\n"
        f"Email: {o.email or '_______________'}"
        f"\n\n"
        f"                                                            ...OPPOSITE PARTY"
    )

    parts.append(
        "COMPLAINT UNDER SECTION 35 OF THE CONSUMER PROTECTION ACT, 2019\n"
        "(Read with the Consumer Protection (Jurisdiction of the District Commission, "
        "the State Commission and the National Commission) Rules, 2021)"
    )

    parts.append("MOST RESPECTFULLY SHOWETH:")

    parts.append("STATEMENT OF FACTS")
    parts.append(narrative.strip())
    parts.append(_grounds_section(intake, sections))

    parts.append("PRAYER:")
    parts.append(
        "In the premises aforesaid and in the interest of justice, it is most "
        "respectfully prayed that this Hon'ble Commission may graciously be pleased to:"
    )
    for i, r in enumerate(_reliefs_block(intake), start=1):
        label = f"({chr(96 + i)})" if i <= 26 else f"({i})"
        parts.append(f"  {label}  {r}")

    parts.append(
        "And pass such other and further order or orders as this Hon'ble Commission "
        "may deem fit and proper in the facts and circumstances of the present case "
        "and in the interest of justice.\n\n"
        "AND FOR THIS ACT OF KINDNESS, THE COMPLAINANT AS IN DUTY BOUND SHALL EVER PRAY."
    )

    parts.append(
        f"Place: {c.city}\n"
        f"Date:  {date.today().strftime('%d.%m.%Y')}\n\n"
        f"_________________________\n"
        f"({c.name})\n"
        f"Complainant"
    )

    parts.append(
        "VERIFICATION:\n\n"
        f"I, {c.name}, the Complainant above named, do hereby solemnly affirm and "
        f"declare that the contents of the above Complaint from Para 1 to Para [__], "
        f"together with the Annexures filed herewith, are true and correct to the best "
        f"of my knowledge and belief. No part thereof is false and nothing material has "
        f"been concealed or suppressed therefrom.\n\n"
        f"Verified at {c.city} on {date.today().strftime('%d.%m.%Y')}.\n\n"
        f"_________________________\n"
        f"({c.name})\n"
        f"Complainant / Deponent"
    )

    return "\n\n".join(parts)


def _write_docx(text: str, path: Path) -> None:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    _COURT_HEADERS = {
        "BEFORE THE HON\u2019BLE",
        "BEFORE THE HON'BLE",
        "BEFORE THE HONS'BLE",
    }
    _SECTION_HEADERS = {
        "MOST RESPECTFULLY SHOWETH:",
        "STATEMENT OF FACTS",
        "GROUNDS:",
        "PRAYER:",
        "VERIFICATION:",
        "VERSUS",
    }

    def _starts_court(block: str) -> bool:
        return block.startswith("BEFORE THE HON")

    def _is_case_no(block: str) -> bool:
        return block.startswith("Consumer Complaint No.")

    def _is_section_hdr(block: str) -> bool:
        return (
            block.strip() in _SECTION_HEADERS
            or block.startswith("COMPLAINT UNDER SECTION")
            or block.startswith("IN THE MATTER OF")
        )

    def _is_party(block: str) -> bool:
        return "...COMPLAINANT" in block or "...OPPOSITE PARTY" in block

    def _add(text_val: str, bold: bool = False, center: bool = False,
             size: int = 12, space_after: int = 6) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text_val)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Times New Roman"

    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            doc.add_paragraph()
            continue

        if _starts_court(block):
            _add(block, bold=True, center=True, size=13, space_after=4)
        elif _is_case_no(block):
            _add(block, bold=True, center=True, size=12, space_after=4)
        elif _is_section_hdr(block):
            _add(block, bold=True, center=False, size=12, space_after=4)
        elif _is_party(block):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            for line in block.split("\n"):
                run = p.add_run(line + "\n")
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                if "...COMPLAINANT" in line or "...OPPOSITE PARTY" in line:
                    run.bold = True
        else:
            _add(block, bold=False, center=False, size=12, space_after=6)

    doc.save(str(path))


def _write_pdf(text: str, path: Path, title: str) -> None:
    doc_pdf = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
        title=title,
    )

    styles = getSampleStyleSheet()

    body = ParagraphStyle(
        "legal_body",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=12,
        leading=18,
        spaceAfter=8,
    )
    heading_center = ParagraphStyle(
        "legal_heading",
        parent=styles["Title"],
        fontName="Times-Bold",
        fontSize=13,
        alignment=1,
        spaceAfter=10,
        spaceBefore=6,
    )
    section_head = ParagraphStyle(
        "section_head",
        parent=styles["Heading2"],
        fontName="Times-Bold",
        fontSize=12,
        alignment=0,
        spaceAfter=6,
        spaceBefore=10,
    )
    party_style = ParagraphStyle(
        "party",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=12,
        leading=16,
        leftIndent=20,
        spaceAfter=4,
    )

    _SECTION_HEADERS = {
        "MOST RESPECTFULLY SHOWETH:",
        "STATEMENT OF FACTS",
        "GROUNDS:",
        "PRAYER:",
        "VERIFICATION:",
        "VERSUS",
    }

    flow = []
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            flow.append(Spacer(1, 0.3 * cm))
            continue

        safe = (
            block
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\n", "<br/>")
        )

        if block.startswith("BEFORE THE HON"):
            flow.append(Paragraph(safe, heading_center))
        elif block.startswith("Consumer Complaint No."):
            flow.append(Paragraph(f"<b>{safe}</b>", heading_center))
        elif (
            block in _SECTION_HEADERS
            or block.startswith("COMPLAINT UNDER SECTION")
            or block.startswith("IN THE MATTER OF")
        ):
            flow.append(Paragraph(f"<b>{safe}</b>", section_head))
        elif "...COMPLAINANT" in block or "...OPPOSITE PARTY" in block:
            flow.append(Paragraph(safe, party_style))
        else:
            flow.append(Paragraph(safe, body))

        flow.append(Spacer(1, 0.15 * cm))

    doc_pdf.build(flow)


def draft_complaint(intake: ComplaintIntake, verdict: JurisdictionVerdict) -> DraftBundle:
    narrative = _polish_narrative(intake)
    sections = _relevant_sections(intake)
    title = (
        f"complaint_{intake.complainant.name.replace(' ', '_').lower()}"
        f"_{int(time.time())}"
    )
    text = _text_complaint(intake, verdict, narrative, sections)
    docx_path = config.OUTPUT_DIR / f"{title}.docx"
    pdf_path = config.OUTPUT_DIR / f"{title}.pdf"
    _write_docx(text, docx_path)
    _write_pdf(text, pdf_path, title=title)
    return DraftBundle(
        title=title,
        polished_narrative=narrative,
        relevant_sections=sections,
        cause_title=_cause_title(intake, verdict),
        docx_path=docx_path,
        pdf_path=pdf_path,
        text=text,
    )
